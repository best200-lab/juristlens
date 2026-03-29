"""
routes/export.py
Export review session as PDF or DOCX
POST /api/juristlens/export
"""

import os
import uuid
import tempfile
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from schemas import ExportRequest, ExportFormat
from services.supabase_service import get_session_messages, verify_session_ownership
from services.claude_service import generate_export_summary

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle

# DOCX generation
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter()


# ─────────────────────────────────────────────
# POST /api/juristlens/export
# ─────────────────────────────────────────────
@router.post("/export")
async def export_review(request: ExportRequest):
    """
    Export the full review session as PDF or DOCX.

    Flow:
    1. Verify lawyer owns this session
    2. Retrieve all Q&A messages from Supabase
    3. Ask Claude to generate a professional summary
    4. Generate PDF or DOCX file
    5. Return file as download
    """

    # Security: verify lawyer owns this session
    owns_session = verify_session_ownership(request.session_id, request.lawyer_id)
    if not owns_session:
        raise HTTPException(status_code=403, detail="Access denied to this session")

    # Get all messages from this session
    messages = get_session_messages(request.session_id)
    if not messages:
        raise HTTPException(status_code=404, detail="No review data found for this session")

    # Generate professional summary via Claude
    summary = generate_export_summary(messages)

    # Generate the file
    try:
        if request.format == ExportFormat.pdf:
            file_path = generate_pdf(messages, summary, request.session_id)
            filename = f"JuristLens_Review_{request.session_id[:8]}.pdf"
            media_type = "application/pdf"
        else:
            file_path = generate_docx(messages, summary, request.session_id)
            filename = f"JuristLens_Review_{request.session_id[:8]}.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return FileResponse(
            path=file_path,
            filename=filename,
            media_type=media_type
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


# ─────────────────────────────────────────────
# Generate PDF Export
# ─────────────────────────────────────────────
def generate_pdf(messages: list, summary: str, session_id: str) -> str:
    """Generate a professional PDF report using ReportLab"""

    # Create temp file
    temp_path = f"/tmp/juristlens_{session_id[:8]}.pdf"

    # Colors matching Jurist Mind brand
    GOLD = colors.HexColor("#c9a84c")
    DARK = colors.HexColor("#0a0a0c")
    DARK_SURFACE = colors.HexColor("#111115")

    doc = SimpleDocTemplate(
        temp_path,
        pagesize=A4,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )

    styles = getSampleStyleSheet()
    story = []

    # ── Header ─────────────────────────────────
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontSize=24,
        textColor=GOLD,
        spaceAfter=6,
        fontName="Helvetica-Bold"
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.grey,
        spaceAfter=20
    )
    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontSize=14,
        textColor=GOLD,
        spaceBefore=20,
        spaceAfter=8,
        fontName="Helvetica-Bold"
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#333333"),
        spaceAfter=8,
        leading=16
    )
    clause_style = ParagraphStyle(
        "Clause",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor("#555555"),
        backColor=colors.HexColor("#FFF8E7"),
        spaceAfter=8,
        leftIndent=20,
        rightIndent=20,
        leading=14,
        borderPadding=(8, 8, 8, 8)
    )

    # Title
    story.append(Paragraph("JuristLens", title_style))
    story.append(Paragraph("Legal Document Review Report", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=GOLD))
    story.append(Spacer(1, 20))

    # Executive Summary
    story.append(Paragraph("Executive Summary", heading_style))
    story.append(Paragraph(summary.replace('\n', '<br/>'), body_style))
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))

    # Individual Q&A Findings
    story.append(Paragraph("Detailed Findings", heading_style))

    for i, msg in enumerate(messages, 1):
        story.append(Paragraph(f"Finding {i}", ParagraphStyle(
            "FindingNum",
            parent=styles["Normal"],
            fontSize=11,
            textColor=GOLD,
            fontName="Helvetica-Bold",
            spaceBefore=15,
            spaceAfter=4
        )))

        story.append(Paragraph(f"<b>Question:</b> {msg['question']}", body_style))
        story.append(Paragraph(f"<b>Answer:</b> {msg['answer']}", body_style))

        if msg.get('clause'):
            story.append(Paragraph(f"Source Clause (Page {msg.get('page_number', 'N/A')}):", ParagraphStyle(
                "ClauseLabel",
                parent=styles["Normal"],
                fontSize=9,
                textColor=colors.grey,
                spaceAfter=2,
                fontName="Helvetica-Bold"
            )))
            story.append(Paragraph(f'"{msg["clause"]}"', clause_style))

        story.append(Spacer(1, 10))

    # Footer note
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "Generated by JuristLens — AI-powered Legal Document Intelligence by Jurist Mind",
        ParagraphStyle("Footer", parent=styles["Normal"], fontSize=8, textColor=colors.grey, alignment=1)
    ))

    doc.build(story)
    return temp_path


# ─────────────────────────────────────────────
# Generate DOCX Export
# ─────────────────────────────────────────────
def generate_docx(messages: list, summary: str, session_id: str) -> str:
    """Generate a professional DOCX report using python-docx"""

    temp_path = f"/tmp/juristlens_{session_id[:8]}.docx"
    doc = Document()

    # ── Page Setup ─────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    GOLD = RGBColor(0xC9, 0xA8, 0x4C)
    DARK = RGBColor(0x33, 0x33, 0x33)
    GREY = RGBColor(0x88, 0x88, 0x88)

    # ── Title ──────────────────────────────────
    title = doc.add_heading("JuristLens", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = GOLD
        run.font.size = Pt(28)
        run.font.bold = True

    subtitle = doc.add_paragraph("Legal Document Review Report")
    subtitle.runs[0].font.color.rgb = GREY
    subtitle.runs[0].font.size = Pt(12)

    doc.add_paragraph("─" * 80)

    # ── Executive Summary ──────────────────────
    summary_heading = doc.add_heading("Executive Summary", 1)
    for run in summary_heading.runs:
        run.font.color.rgb = GOLD

    summary_para = doc.add_paragraph(summary)
    summary_para.runs[0].font.size = Pt(10)

    doc.add_paragraph("─" * 80)

    # ── Detailed Findings ──────────────────────
    findings_heading = doc.add_heading("Detailed Findings", 1)
    for run in findings_heading.runs:
        run.font.color.rgb = GOLD

    for i, msg in enumerate(messages, 1):
        # Finding number
        finding_heading = doc.add_heading(f"Finding {i}", 2)
        for run in finding_heading.runs:
            run.font.color.rgb = GOLD
            run.font.size = Pt(12)

        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run("Question: ")
        q_run.bold = True
        q_run.font.color.rgb = DARK
        q_para.add_run(msg['question']).font.color.rgb = DARK

        # Answer
        a_para = doc.add_paragraph()
        a_run = a_para.add_run("Answer: ")
        a_run.bold = True
        a_run.font.color.rgb = DARK
        a_para.add_run(msg['answer']).font.color.rgb = DARK

        # Source clause
        if msg.get('clause'):
            page_label = doc.add_paragraph()
            page_run = page_label.add_run(f"Source Clause (Page {msg.get('page_number', 'N/A')}):")
            page_run.bold = True
            page_run.font.size = Pt(9)
            page_run.font.color.rgb = GREY

            clause_para = doc.add_paragraph(f'"{msg["clause"]}"')
            clause_para.runs[0].font.size = Pt(9)
            clause_para.runs[0].font.italic = True
            clause_para.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Add paragraph border styling
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = clause_para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '4')
                border.set(qn('w:color'), 'C9A84C')
                pBdr.append(border)
            pPr.append(pBdr)

        doc.add_paragraph()  # Spacing between findings

    # ── Footer ─────────────────────────────────
    doc.add_paragraph("─" * 80)
    footer = doc.add_paragraph(
        "Generated by JuristLens — AI-powered Legal Document Intelligence by Jurist Mind"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = GREY

    doc.save(temp_path)
    return temp_path